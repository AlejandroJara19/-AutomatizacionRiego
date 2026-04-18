# --- INTERFAZ Y GRÁFICOS ---
import streamlit as st
import plotly.express as px
import plotly.io as pio
import folium
from streamlit_folium import st_folium

# --- MANEJO DE DATOS Y CÁLCULOS ---
import pandas as pd
import numpy as np
import math
import json
import datetime

# --- GEOPROCESAMIENTO (Sin GDAL pesado) ---
import rasterio
from rasterio.mask import mask
from rasterio.io import MemoryFile
import geopandas as gpd
from shapely.geometry import Point, mapping
from shapely.ops import transform

# --- UTILIDADES DEL SISTEMA ---
import os
import shutil
import tempfile
import zipfile
import io
import requests

# --- GENERACIÓN DE INFORMES (ADR) ---
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# --- PRUEBAS (ADR) ---
import streamlit as st
import pandas as pd
import io

# --- CACHÉ / RED ---
REQUEST_TIMEOUT = 30

@st.cache_data(show_spinner=False)
def fetch_nasa_data(lat, lon, fecha_inicio_str, fecha_fin_str):
    url_nasa = (
        f"https://power.larc.nasa.gov/api/temporal/daily/point"
        f"?parameters=PRECTOTCORR,T2M_MAX,T2M_MIN,EVPTRNS"
        f"&community=AG&longitude={lon}&latitude={lat}"
        f"&start={fecha_inicio_str}&end={fecha_fin_str}&format=JSON"
    )
    response = requests.get(url_nasa, timeout=REQUEST_TIMEOUT)
    response.raise_for_status()
    data = response.json()
    params = data['properties']['parameter']
    df_nasa = pd.DataFrame({
        'Fecha': list(params['PRECTOTCORR'].keys()),
        'Precipitacion': list(params['PRECTOTCORR'].values()),
        'T_Max': list(params['T2M_MAX'].values()),
        'T_Min': list(params['T2M_MIN'].values()),
        'Evaporacion': list(params['EVPTRNS'].values())
    })
    df_nasa.replace(-999.0, np.nan, inplace=True)
    df_nasa.fillna(0, inplace=True)
    df_nasa['Fecha'] = pd.to_datetime(df_nasa['Fecha'], format='%Y%m%d')
    return df_nasa

def calcular_ret_vectorizado(df_nasa, lat_input):
    df = df_nasa.copy()
    df['DOY'] = df['Fecha'].dt.dayofyear
    lat_rad = math.radians(lat_input)

    t_max = df['T_Max'].to_numpy(dtype=float)
    t_min = df['T_Min'].to_numpy(dtype=float)
    doy = df['DOY'].to_numpy(dtype=float)
    t_mean = (t_max + t_min) / 2.0

    dr = 1.0 + 0.033 * np.cos(2.0 * np.pi * doy / 365.0)
    delta = 0.409 * np.sin((2.0 * np.pi * doy / 365.0) - 1.39)
    ws_arg = -np.tan(lat_rad) * np.tan(delta)
    ws_arg = np.clip(ws_arg, -1.0, 1.0)
    ws = np.arccos(ws_arg)

    ra = (24.0 * 60.0 / np.pi) * 0.0820 * dr * (
        ws * np.sin(lat_rad) * np.sin(delta) +
        np.cos(lat_rad) * np.cos(delta) * np.sin(ws)
    )
    ra_mm = ra * 0.408
    delta_t = np.maximum(t_max - t_min, 0.0)

    ret = np.where(
        t_max > t_min,
        0.0023 * ra_mm * (t_mean + 17.8) * np.sqrt(delta_t),
        0.0
    )
    df['RET'] = ret
    return df

def preparar_base_nasa(lat_input, lon_input, fecha_inicio, fecha_fin):
    f_inicio_nasa = fecha_inicio.strftime("%Y%m%d")
    f_fin_nasa = fecha_fin.strftime("%Y%m%d")
    df_nasa = fetch_nasa_data(lat_input, lon_input, f_inicio_nasa, f_fin_nasa)
    df_nasa = calcular_ret_vectorizado(df_nasa, lat_input)
    return df_nasa

def agregar_decadas(df_base_diario):
    df = df_base_diario.copy()
    df['Año'] = df['Fecha'].dt.year
    df['Mes'] = df['Fecha'].dt.month
    df['Día'] = df['Fecha'].dt.day
    df['Decada_Mes'] = np.select(
        [df['Día'] <= 10, df['Día'] <= 20],
        [1, 2],
        default=3
    )
    df['Decada_Año'] = (df['Mes'] - 1) * 3 + df['Decada_Mes']
    return df

def crear_memoria_hidrologia(datos_clima, coordenadas, df_simulacion=None, tipo_almacenamiento="No definido", vol_max=0):
    doc = Document()
    
    # Encabezado ADR
    titulo = doc.add_heading('Anexo 3: Memoria de Cálculo Hidrología y Climatología', 0)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # --- Sección 1 a 3 (Se mantienen intactas) ---
    doc.add_heading('1. Metodología Climatológica', level=1)
    p = doc.add_paragraph()
    p.add_run('Para el dimensionamiento de los Sistemas de Riego Individuales o Comunitarios (SRIC), ').bold = True
    p.add_run('se ha utilizado información satelital de la plataforma WaPOR v3. La serie de datos analizada comprende un periodo de 10 años, garantizando la representatividad climática de la zona.')

    doc.add_heading('2. Análisis de Precipitación Decadal', level=2)
    doc.add_paragraph("Se define la 'década' como un periodo de exactamente 10 días (o el remanente del mes). Este intervalo permite capturar la variabilidad de la humedad en el suelo de forma más precisa que un promedio mensual.")

    doc.add_heading('3. Cálculo de Probabilidad de Excedencia (P75)', level=2)
    doc.add_paragraph("Se aplicó el método de Critchley y Siegert (1996) para determinar la precipitación persistente. La fórmula aplicada es:")
    doc.add_paragraph("P % = [(m - 0.375) / (N + 0.25)] * 100", style='Intense Quote')
    doc.add_paragraph("Donde 'm' es el orden de la serie y 'N' el número total de observaciones. Se asume el P75 como el valor de diseño para asegurar el suministro en años secos.")

    # --- NUEVA SECCIÓN 4: Infraestructura de Almacenamiento ---
    doc.add_heading('4. Dimensionamiento del Sistema de Almacenamiento', level=1)
    p_almacenamiento = doc.add_paragraph(f"Para el presente proyecto SRIC, se ha seleccionado la alternativa constructiva de: ")
    p_almacenamiento.add_run(f"{tipo_almacenamiento}").bold = True
    p_almacenamiento.add_run(f", con una capacidad máxima proyectada (NAME) de {vol_max:.2f} m³.")

    # Justificación técnica según el tipo de reservorio elegido
    if "Excavado" in tipo_almacenamiento:
        doc.add_paragraph("Dado que se trata de un reservorio excavado con geometría irregular, el espejo de agua y el volumen almacenado varían en función de la altura de la lámina de agua. Para el cálculo riguroso de las pérdidas por evaporación, se empleó una curva Cota-Área-Volumen derivada de la batimetría del sitio, ajustada numéricamente.")
    else:
        doc.add_paragraph("Al tratarse de un tanque australiano, se asume un área de espejo de agua constante (cilindro) para el cálculo de los aportes por precipitación directa y las salidas por evaporación libre.")

    # --- NUEVA SECCIÓN 5: Balance Hídrico (Tabla 2 exigida por la ADR) ---
    doc.add_heading('5. Funcionamiento del vaso de almacenamiento', level=1)
    doc.add_paragraph("La simulación del tránsito del embalse se realizó a nivel decadal, evaluando las entradas frente a las salidas. A continuación, se presenta la Tabla 2 con el balance hídrico del sistema:")

    if df_simulacion is not None and not df_simulacion.empty:
        # Título de la tabla
        doc.add_paragraph("Tabla 2. Funcionamiento del vaso de almacenamiento", style='Caption')
        
        # Crear la tabla en Word con 8 columnas (Estándar ADR)
        columnas_tabla = ['AÑO', 'MES', 'DÉCADA', 'VOL. INICIAL (m3)', 'ENTRADAS (+)', 'SALIDAS (-)', 'VOL. FINAL (m3)', 'EXCEDENTE (m3)']
        tabla = doc.add_table(rows=1, cols=len(columnas_tabla))
        tabla.style = 'Table Grid'
        
        # Agregar encabezados
        hdr_cells = tabla.rows[0].cells
        for i, col_name in enumerate(columnas_tabla):
            hdr_cells[i].text = col_name

        # Llenar datos: Mostrar solo un año tipo (primeras 36 décadas) para que el Word no quede de 50 páginas
        # Se suman las variables de la Pestaña 3 para condensarlas en "Entradas" y "Salidas"
        for index, row in df_simulacion.head(36).iterrows():
            row_cells = tabla.add_row().cells
            
            # Cálculos internos para unificar las entradas y salidas de tu df_simulacion
            entradas_totales = row.get('Entrada Concesion (m3)', 0) + row.get('Entrada Lluvia (m3)', 0) + row.get('Entrada Escorrentia (m3)', 0)
            salidas_totales = row.get('Salida Riego (m3)', 0) + row.get('Salida Evaporación (m3)', 0) + row.get('Salida Infiltración (m3)', 0)
            
            # Asignación a las celdas
            row_cells[0].text = str(int(row.get('Año', 0)))
            row_cells[1].text = "N/A" # O el mes si lo tienes en tu DF
            row_cells[2].text = str(int(row.get('Decada', 0)))
            row_cells[3].text = f"{row.get('Volumen Inicial (m3)', 0):.2f}"
            row_cells[4].text = f"{entradas_totales:.2f}"
            row_cells[5].text = f"{salidas_totales:.2f}"
            row_cells[6].text = f"{row.get('Volumen Final (m3)', 0):.2f}"
            row_cells[7].text = f"{row.get('Volumen Derramado (m3)', 0):.2f}"
            
        p_nota = doc.add_paragraph()
        p_nota.add_run("*Nota: Por extensión, se presentan los resultados correspondientes al primer año de simulación (36 décadas). El balance histórico completo reposa en los archivos digitales del proyecto.").italic = True
    else:
        doc.add_paragraph("⚠️ Error: No se encontraron datos de simulación. Ejecute la Pestaña 3 primero.")

    return doc

def crear_memoria_demandas(datos_cultivo):
    doc = Document()
    doc.add_heading('Memoria de Cálculo: Disponibilidad y Demandas hídricas', 0)
    
    doc.add_heading('1. Demanda Hídrica del Cultivo (ETc)', level=1)
    doc.add_paragraph(
        "El cálculo de la demanda se basa en la interacción entre la Evapotranspiración de Referencia (ET0) "
        "y el Coeficiente de Cultivo (Kc) específico para cada etapa fenológica."
    )
    
    doc.add_paragraph("ETc = ET0 * Kc", style='Intense Quote')
    
    doc.add_heading('2. Requerimiento de Riego', level=2)
    doc.add_paragraph(
        "Considerando la precipitación efectiva calculada en el Anexo 3, el requerimiento neto (Rn) se define como:"
    )
    doc.add_paragraph("Rn = ETc - P_efectiva", style='Intense Quote')
    
    return doc

# Inicialización de variables de estado (Session State)
# Coloca esto en la parte superior de tu app.py, fuera de cualquier pestaña
if 'zip_buffer' not in st.session_state:
    st.session_state.zip_buffer = None
if 'process_complete' not in st.session_state:
    st.session_state.process_complete = False

# --- CONFIGURACIÓN DE PÁGINA ---
st.set_page_config(page_title="HidroApp ADR", page_icon="💧", layout="wide")
st.title("💧 HidroApp: Análisis y Descarga de Datos Climáticos")

def extraer_fecha_de_nombre(nombre_archivo):
    """
    Extrae fecha desde nombres con formatos:
    - YYYYMMDD
    - YYYY-MM-DD
    - YYYY_MM_DD
    """
    import re
    from datetime import datetime

    match = re.search(r'(\d{4})[-_]?(\d{2})[-_]?(\d{2})', nombre_archivo)
    if not match:
        return None

    year, month, day = match.groups()
    try:
        return datetime(int(year), int(month), int(day))
    except ValueError:
        return None

def procesar_zip_wapor(archivo_zip, lon, lat, nombre_columna):
    """
    Lee un archivo ZIP con rasters de WaPOR en memoria y extrae el valor para una coordenada.
    Retorna un DataFrame con 'Fecha' y la columna de la variable (ej. 'Precipitacion').
    """
    resultados = []
    archivos_tif = 0
    archivos_sin_fecha = 0
    archivos_con_error = 0

    with zipfile.ZipFile(archivo_zip) as z:
        for filename in z.namelist():
            if not filename.lower().endswith((".tif", ".tiff")):
                continue

            archivos_tif += 1

            fecha_extraida = extraer_fecha_de_nombre(os.path.basename(filename))
            if fecha_extraida is None:
                archivos_sin_fecha += 1
                continue

            fecha_obj = pd.to_datetime(fecha_extraida)

            try:
                with z.open(filename) as f:
                    with MemoryFile(f.read()) as memfile:
                        with memfile.open() as src:
                            muestra = list(src.sample([(lon, lat)]))
                            if not muestra or len(muestra[0]) == 0:
                                valor = np.nan
                            else:
                                valor = float(muestra[0][0])

                            # Limpieza de valores anómalos o "No Data" típicos de raster
                            if np.isnan(valor) or valor < -99 or valor > 10000:
                                valor = 0.0

                            resultados.append({'Fecha': fecha_obj, nombre_columna: valor})
            except Exception:
                archivos_con_error += 1
                continue

    # Retorno consistente para evitar KeyError al hacer merge por 'Fecha'
    if resultados:
        df_resultado = pd.DataFrame(resultados)
        df_resultado = df_resultado.sort_values('Fecha').drop_duplicates(subset=['Fecha']).reset_index(drop=True)
    else:
        df_resultado = pd.DataFrame(columns=['Fecha', nombre_columna])

    return df_resultado, {
        'archivos_tif': archivos_tif,
        'archivos_validos': len(resultados),
        'archivos_sin_fecha': archivos_sin_fecha,
        'archivos_con_error': archivos_con_error
    }

# --- CREACIÓN DE PESTAÑAS ---
tab1, tab2, tab3, tab4 = st.tabs([
    "📊 Datos Agroclimáticos", 
    "💧 Balance Hídrico", 
    "📈 Volúmenes de Riego", 
    "⚙️ Generación WaPOR"
])
# =====================================================================
# --- PESTAÑA 1: ANÁLISIS DEL EXCEL / CSV --- (Sin cambios)
# =====================================================================
with tab1:
    st.markdown("Descarga la serie de la NASA o extrae información climática de repositorios Raster (WaPOR v3) para analizar la **Precipitación Confiable al 75%** y promedios decadales.")
    
    fuente_datos = st.radio("📡 Seleccione la fuente de datos:", ["NASA POWER (API Online)", "WaPOR v3 (Archivos Raster .ZIP)"], horizontal=True)
    
    col1, col2 = st.columns(2)
    with col1:
        lat_input = st.number_input("Latitud (Ej: 4.7160)", value=6.326512, format="%.6f", key="lat_t1")
    with col2:
        lon_input = st.number_input("Longitud (Ej: -74.2160)", value=-73.609413, format="%.6f", key="lon_t1")

    # Variable global de la pestaña que almacenará la serie diaria sin importar el origen
    df_base_diario = None 

    # ==========================================
    # RAMA 1: DESCARGA DESDE NASA POWER
    # ==========================================
    if fuente_datos == "NASA POWER (API Online)":
        col3, col4 = st.columns(2)
        with col3:
            fecha_inicio = st.date_input("Fecha Inicio", pd.to_datetime("2018-01-01"), key="fi_nasa")
        with col4:
            fecha_fin = st.date_input("Fecha Fin", pd.to_datetime("2025-12-31"), key="ff_nasa")

        if st.button("Obtener Datos NASA y Calcular Balance", type="primary"):
            with st.spinner('Consultando a la NASA y calculando variables hídricas... 🚀'):
                try:
                    df_base_diario = preparar_base_nasa(lat_input, lon_input, fecha_inicio, fecha_fin)
                    st.success("✅ ¡Datos de NASA descargados con éxito!")
                except Exception as e:
                    st.error(f"Error técnico con NASA: {e}")

    # ==========================================
    # RAMA 2: EXTRACCIÓN WAPOR v3 (.ZIP)
    # ==========================================
    elif fuente_datos == "WaPOR v3 (Archivos Raster .ZIP)":
        st.info("Sube los archivos .zip que contienen los TIFs de cada variable. El sistema cruzará la información con la coordenada ingresada.")
        
        col_w1, col_w2, col_w3 = st.columns(3)
        with col_w1:
            zip_precip = st.file_uploader("ZIP Precipitación", type="zip")
        with col_w2:
            zip_evap = st.file_uploader("ZIP Evaporación", type="zip")
        with col_w3:
            zip_ret = st.file_uploader("ZIP Evapotranspiración (RET)", type="zip")
            
        if st.button("Procesar Datos WaPOR y Calcular Balance", type="primary"):
            if zip_precip and zip_evap and zip_ret:
                with st.spinner('Extrayendo píxeles de los archivos Raster. Esto puede tardar unos segundos... 🛰️'):
                    try:
                        df_p, rep_p = procesar_zip_wapor(zip_precip, lon_input, lat_input, 'Precipitacion')
                        df_e, rep_e = procesar_zip_wapor(zip_evap, lon_input, lat_input, 'Evaporacion')
                        df_r, rep_r = procesar_zip_wapor(zip_ret, lon_input, lat_input, 'RET')

                        # Validaciones explícitas para evitar KeyError: 'Fecha'
                        faltantes = []
                        if df_p.empty:
                            faltantes.append(
                                f"Precipitación: 0 fechas válidas (TIF: {rep_p['archivos_tif']}, sin fecha: {rep_p['archivos_sin_fecha']}, error lectura: {rep_p['archivos_con_error']})"
                            )
                        if df_e.empty:
                            faltantes.append(
                                f"Evaporación: 0 fechas válidas (TIF: {rep_e['archivos_tif']}, sin fecha: {rep_e['archivos_sin_fecha']}, error lectura: {rep_e['archivos_con_error']})"
                            )
                        if df_r.empty:
                            faltantes.append(
                                f"RET: 0 fechas válidas (TIF: {rep_r['archivos_tif']}, sin fecha: {rep_r['archivos_sin_fecha']}, error lectura: {rep_r['archivos_con_error']})"
                            )

                        if faltantes:
                            st.error("No se pudieron extraer fechas/datos válidos de uno o más ZIP.")
                            for msg in faltantes:
                                st.warning(msg)
                        else:
                            # Unimos las 3 variables basadas en la fecha
                            df_clima = pd.merge(df_p, df_e, on='Fecha', how='inner')
                            df_clima = pd.merge(df_clima, df_r, on='Fecha', how='inner')

                            if df_clima.empty:
                                st.warning("No se encontraron fechas coincidentes entre los tres ZIP cargados.")
                            else:
                                df_base_diario = df_clima # Pasamos la estafeta al bloque común
                                st.success("✅ ¡Datos Raster procesados con éxito!")
                    except Exception as e:
                        st.error(f"Error procesando los archivos Raster: {e}")
            else:
                st.warning("Por favor, sube los 3 archivos ZIP para continuar.")


    # ==========================================
    # BLOQUE COMÚN: PROCESAMIENTO DECADAL Y GRÁFICAS
    # (Se ejecuta si df_base_diario fue llenado por NASA o por WaPOR)
    # ==========================================
    if df_base_diario is not None and not df_base_diario.empty:
        df_base_diario = agregar_decadas(df_base_diario)
        
        # 1. Sumar los valores por década para cada año individual
        df_decadal_anual = df_base_diario.groupby(['Año', 'Decada_Año'])[['Precipitacion', 'Evaporacion', 'RET']].sum().reset_index()
        
        # 2. Promediar Evaporación y RET a lo largo del periodo
        df_promedio_decadal = df_decadal_anual.groupby('Decada_Año')[['Evaporacion', 'RET']].mean().reset_index()
        
        # 3. Precipitación al 75% (Percentil 25)
        def calcular_prob_75(serie):
            if len(serie) == 0: return 0.0
            return np.percentile(serie, 25)

        precip_75_serie = df_decadal_anual.groupby('Decada_Año')['Precipitacion'].apply(calcular_prob_75).reset_index()
        precip_75_serie.rename(columns={'Precipitacion': 'Prec_75%'}, inplace=True)
        
        # Unir la precipitación calculada con los promedios
        df_promedio_decadal = pd.merge(df_promedio_decadal, precip_75_serie, on='Decada_Año')
        
        # --- EVIDENCIA DEL ANÁLISIS DE PROBABILIDAD (BLOM) ---
        st.subheader("🌧️ Análisis de Probabilidad (Precipitación ordenada)")
        try:
            df_prob_global = pd.DataFrame()
            for decada in range(1, 37):
                serie_decada = df_decadal_anual[df_decadal_anual['Decada_Año'] == decada]['Precipitacion'].values
                df_prob_global[f'D{decada}'] = np.sort(serie_decada)[::-1]
                
            N_anios = len(df_prob_global)
            m_arr = np.arange(1, N_anios + 1)
            prob_arr = ((m_arr - 0.375) / (N_anios + 0.25)) * 100
            
            df_prob_global.insert(0, 'Probabilidad Blom (%)', prob_arr)
            df_prob_global.insert(0, 'Orden (m)', m_arr)
            
            with st.expander("Ver matriz completa de ordenamiento y probabilidades"):
                st.dataframe(df_prob_global.round(2).style.format("{:.2f}", subset=['Probabilidad Blom (%)'] + [f'D{i}' for i in range(1, 37)]))
        except Exception as e:
            st.warning("No se pudo generar la tabla de visualización de probabilidad.")

        # --- GRÁFICA INTERACTIVA ---
        st.subheader("📈 Comportamiento Hídrico Decadal")
        df_melted = df_promedio_decadal.melt(
            id_vars=['Decada_Año'], value_vars=['Prec_75%', 'Evaporacion', 'RET'],
            var_name='Variable', value_name='Volumen (mm/década)'
        )
        
        import plotly.express as px
        fig = px.line(
            df_melted, x='Decada_Año', y='Volumen (mm/década)', color='Variable', 
            markers=True, color_discrete_map={'Prec_75%': '#1f77b4', 'Evaporacion': '#ff7f0e', 'RET': '#d62728'},
            labels={'Decada_Año': 'Década del Año (1 al 36)'}
        )
        fig.update_layout(xaxis=dict(tickmode='linear', tick0=1, dtick=1), hovermode="x unified")
        st.plotly_chart(fig)
        
        # --- GENERAR IMAGEN ESTÁTICA SEGURA (Matplotlib) ---
        try:
            import matplotlib.pyplot as plt
            import io
            fig_static, ax = plt.subplots(figsize=(8, 4))
            ax.plot(df_promedio_decadal['Decada_Año'], df_promedio_decadal['Prec_75%'], label='Prec_75%', color='#1f77b4', marker='o', markersize=4)
            ax.plot(df_promedio_decadal['Decada_Año'], df_promedio_decadal['Evaporacion'], label='Evaporacion', color='#ff7f0e', marker='o', markersize=4)
            ax.plot(df_promedio_decadal['Decada_Año'], df_promedio_decadal['RET'], label='RET', color='#d62728', marker='o', markersize=4)
            ax.set_title('Comportamiento Hídrico Decadal', fontsize=12)
            ax.set_xlabel('Década del Año (1 al 36)', fontsize=10)
            ax.set_ylabel('Volumen (mm/década)', fontsize=10)
            ax.legend()
            ax.grid(True, linestyle='--', alpha=0.7)
            
            img_buffer = io.BytesIO()
            fig_static.savefig(img_buffer, format='png', bbox_inches='tight', dpi=150)
            img_buffer.seek(0)
            st.session_state['imagen_clima_bytes'] = img_buffer.getvalue()
            plt.close(fig_static) 
        except Exception as e:
            st.session_state['imagen_clima_bytes'] = None

        # --- GUARDADO EN VARIABLES DE SESIÓN PARA LA PESTAÑA 2 ---
        st.session_state['latitud'] = lat_input
        st.session_state['longitud'] = lon_input
        st.session_state['df_promedio'] = df_promedio_decadal
        st.session_state['df_base_diario_tab1'] = df_base_diario.copy()
        
        # --- DESCARGA DE DATOS ---
        st.subheader("📥 Descarga de Resultados")
        col_down1, col_down2 = st.columns(2)
        
        with col_down1:
            with st.expander("Ver tabla de resultados decadales"):
                st.dataframe(df_promedio_decadal.round(2).style.format("{:.2f}"))
            csv_promedio = df_promedio_decadal.round(2).to_csv(index=False, sep=";")
            st.download_button("📥 Descargar Datos Decadales (CSV)", data=csv_promedio, file_name=f"Balance_Decadal_{lat_input}_{lon_input}.csv", mime="text/csv", key="btn_down_promedios")

        with col_down2:
            with st.expander("Ver datos diarios originales (NASA o WaPOR)"):
                st.dataframe(df_base_diario[['Fecha', 'Año', 'Mes', 'Día', 'Precipitacion', 'Evaporacion', 'RET']].head(100))
                st.caption("Mostrando los primeros 100 registros.")
            csv_diario = df_base_diario.round(2).to_csv(index=False, sep=";")
            st.download_button("📥 Descargar Serie Diaria (CSV)", data=csv_diario, file_name=f"Serie_Diaria_{lat_input}_{lon_input}.csv", mime="text/csv", key="btn_down_diario")
# =====================================================================
# --- PESTAÑA 2: DESCARGA AUTOMÁTICA NASA POWER --- (Sin cambios)
# =====================================================================

with tab2:
        st.markdown("### Balance Hídrico y Diseño Hidráulico")
        st.markdown("Determinación paso a paso de las necesidades netas, brutas y el dimensionamiento de caudales de diseño por sector.")
        
        # --- SECCIÓN 1: DATOS CLIMÁTICOS ---
        st.subheader("1. Ubicación y Periodo")
        fuente_clima_t2 = st.radio(
            "Fuente climática para Pestaña 2:",
            ["Usar datos procesados en Pestaña 1", "NASA POWER (API Online)"],
            horizontal=True,
            key="fuente_clima_t2"
        )

        col1, col2 = st.columns(2)
        with col1:
            lat_input = st.number_input("Latitud", value=6.326512, format="%.6f", key="lat_nasa_t2")
        with col2:
            lon_input = st.number_input("Longitud", value=-73.609413, format="%.6f", key="lon_nasa_t2")

        col3, col4 = st.columns(2)
        with col3:
            fecha_inicio = st.date_input("Fecha Inicio", pd.to_datetime("2018-01-01"), key="fi_nasa_t2")
        with col4:
            fecha_fin = st.date_input("Fecha Fin", pd.to_datetime("2025-12-31"), key="ff_nasa_t2")

        # --- SECCIÓN 2: PARÁMETROS AGRONÓMICOS ---
        st.subheader("2. Parámetros Agronómicos y Fenología")
        ca1, ca2, ca3 = st.columns(3)
        with ca1:
            area_total = st.number_input("Área Total (Ha)", value=0.50, step=0.1, min_value=0.01, key="area_tot")
        with ca2:
            num_sectores = st.number_input("Número de Sectores", value=1, min_value=1, step=1, key="num_sect")
        with ca3:
            decada_inicio = st.number_input("Década de Inicio (1-36)", min_value=1, max_value=36, value=1, key="dec_ini")
            
        siembra_escalonada = st.checkbox("¿Aplicar siembra escalonada entre sectores?", value=True, key="check_esc")
        paso_escalonamiento = st.number_input("Décadas de espera entre siembra", min_value=0, value=1, key="paso_esc") if siembra_escalonada else 0

        st.markdown("**Selección de Cultivo (Seguridad Alimentaria - FAO 56)**")
        
        # Base de datos de cultivos (Duraciones en décadas de 10 días) - Sin Plátano
        base_cultivos = {
            "Maíz (Grano Seco)": {"kc_ini": 0.30, "kc_mid": 1.20, "kc_end": 0.50, "L_ini": 3, "L_dev": 4, "L_mid": 4, "L_late": 3},
            "Maíz (Dulce/Húmedo)": {"kc_ini": 0.30, "kc_mid": 1.15, "kc_end": 1.05, "L_ini": 2, "L_dev": 3, "L_mid": 3, "L_late": 1},
            "Frijol Seco": {"kc_ini": 0.40, "kc_mid": 1.15, "kc_end": 0.35, "L_ini": 2, "L_dev": 3, "L_mid": 4, "L_late": 2},
            "Yuca (Cassava)": {"kc_ini": 0.30, "kc_mid": 1.10, "kc_end": 0.50, "L_ini": 2, "L_dev": 4, "L_mid": 15, "L_late": 6},
            "Ñame (Yam)": {"kc_ini": 0.30, "kc_mid": 1.10, "kc_end": 0.60, "L_ini": 6, "L_dev": 8, "L_mid": 12, "L_late": 4},
            "Personalizado": {"kc_ini": 0.40, "kc_mid": 1.10, "kc_end": 0.60, "L_ini": 3, "L_dev": 4, "L_mid": 4, "L_late": 3}
        }

        cultivo_seleccionado = st.selectbox("Seleccione el cultivo a establecer:", list(base_cultivos.keys()))
        datos_c = base_cultivos[cultivo_seleccionado]

        # Mostrar inputs editables por si el usuario quiere ajustar la base de datos localmente
        ck1, ck2, ck3, ck4 = st.columns(4)
        with ck1:
            kc_ini = st.number_input("Kc Inicial", value=datos_c["kc_ini"], step=0.05, key="kc_i")
            L_ini = st.number_input("Dur. Inicial (Décadas)", value=datos_c["L_ini"], min_value=1, key="l_i")
        with ck2:
            kc_mid = st.number_input("Kc Medio", value=datos_c["kc_mid"], step=0.05, key="kc_m")
            L_dev = st.number_input("Dur. Desarrollo", value=datos_c["L_dev"], min_value=1, key="l_d")
        with ck3:
            kc_end = st.number_input("Kc Final", value=datos_c["kc_end"], step=0.05, key="kc_f")
            L_mid = st.number_input("Dur. Media", value=datos_c["L_mid"], min_value=1, key="l_m")
        with ck4:
            st.write("") # Espaciador
            st.write("")
            L_late = st.number_input("Dur. Final (Maduración)", value=datos_c["L_late"], min_value=1, key="l_l")

        duracion_total = int(L_ini + L_dev + L_mid + L_late)
        
        # --- CONSTRUCCIÓN DE LA CURVA KC (Interpolación Lineal FAO) ---
        curva_kc = []
        # 1. Fase Inicial (Constante)
        curva_kc.extend([kc_ini] * int(L_ini))
        # 2. Fase de Desarrollo (Interpolación de kc_ini a kc_mid)
        if L_dev > 0:
            paso_dev = (kc_mid - kc_ini) / L_dev
            curva_kc.extend([kc_ini + paso_dev * (i + 1) for i in range(int(L_dev))])
        # 3. Fase Media (Constante)
        curva_kc.extend([kc_mid] * int(L_mid))
        # 4. Fase Final (Interpolación de kc_mid a kc_end)
        if L_late > 0:
            paso_late = (kc_end - kc_mid) / L_late
            curva_kc.extend([kc_mid + paso_late * (i + 1) for i in range(int(L_late))])

        # Visualización de la Curva
        import plotly.graph_objects as go
        fig_kc = go.Figure()
        fig_kc.add_trace(go.Scatter(x=list(range(1, duracion_total + 1)), y=curva_kc, mode='lines+markers', name='Curva Kc', line=dict(color='DarkGreen', width=3)))
        fig_kc.update_layout(title=f"Curva Fenológica del {cultivo_seleccionado} (Ciclo: {duracion_total*10} días)", xaxis_title="Décadas (10 días)", yaxis_title="Coeficiente de Cultivo (Kc)", height=350, margin=dict(l=0, r=0, t=40, b=0))
        st.plotly_chart(fig_kc, use_container_width=True)

        sembrar_multiple = st.checkbox("Habilitar múltiples ciclos de producción", value=True, key="check_mult") if duracion_total < 36 else False
        descanso = st.number_input("Décadas de descanso", min_value=0, value=1, key="descanso") if sembrar_multiple else 0
        
        # --- SECCIÓN 3: CONFIGURACIÓN DEL SISTEMA DE RIEGO ---
        st.subheader("3. Configuración del Sistema de Riego")
        
        # Nuevo selector para elegir el tipo de riego
        tipo_riego = st.radio("Tipo de Riego", options=["Riego por goteo", "Riego por aspersión"], horizontal=True, key="tipo_riego")
        
        if tipo_riego == "Riego por goteo":
            cs1, cs2, cs3, cs4 = st.columns(4)
            with cs1: dist_emisores = st.number_input("Dist. Emisores (m)", value=0.20, step=0.05, min_value=0.01, key="dist_e")
            with cs2: dist_laterales = st.number_input("Dist. Laterales (m)", value=1.20, step=0.05, min_value=0.01, key="dist_l")
            with cs3: emisores_planta = st.number_input("Emisores / Planta", value=2, min_value=1, key="em_pl")
            with cs4: caudal_emisor_lh = st.number_input("Caudal Gotero (L/h)", value=1.00, step=0.1, min_value=0.01, key="q_got")
        else:
            # Inputs exclusivos para aspersión
            cs1, cs2 = st.columns(2)
            with cs1: num_aspersores_ha = st.number_input("Número de aspersores por ha", value=100, min_value=1, step=1, key="num_asp_ha")
            with cs2: caudal_emisor_lh = st.number_input("Caudal del Aspersor (L/h)", value=500.0, step=10.0, min_value=1.0, key="q_asp")

        # Inputs de eficiencia
        ce1, ce2, ce3, ce4, ce5 = st.columns(5)
        with ce1: area_sombreada = st.number_input("% Sombreado", value=65.0, step=5.0, min_value=0.0, max_value=100.0, key="a_som")
        with ce2: pct_sustrato = st.number_input("% Sustrato", value=100.0, step=5.0, min_value=0.0, max_value=100.0, key="p_sus")
        with ce3: ef_cond = st.number_input("% Ef. Cond", value=98.0, step=1.0, min_value=0.0, max_value=100.0, key="ef_c")
        with ce4: ef_dist = st.number_input("% Ef. Dist", value=98.0, step=1.0, min_value=0.0, max_value=100.0, key="ef_d")
        with ce5: ef_rieg = st.number_input("% Ef. Riego", value=90.0, step=1.0, min_value=0.0, max_value=100.0
               
                
# =====================================================================
# --- PESTAÑA 3: FUNCIONAMIENTO RESERVORIO
# =====================================================================

with tab3:
    st.markdown("### Simulación Cronológica del Reservorio")
    st.markdown("Tránsito del embalse frente a la serie climática histórica para evaluar el riesgo de déficit.")

    # 1. Selección de Infraestructura Principal
    tipo_almacenamiento = st.radio(
        "Seleccione el tipo de estructura de almacenamiento:",
        ["Opción 1: Tanque Australiano (Cilíndrico)", "Opción 3: Reservorio Excavado (Vaso Irregular)"],
        help="La Opción 3 utiliza curvas de nivel para un cálculo más preciso en terrenos irregulares."
    )

    # Variables de inicialización
    es_excavado = False
    vol_max_sistema = 0.0
    area_fija_espejo = 0.0

    if tipo_almacenamiento == "Opción 1: Tanque Australiano (Cilíndrico)":
        st.subheader("1. Dimensiones del Tanque Australiano")
        col1, col2, col3 = st.columns(3)
        with col1:
            radio_tanque = st.number_input("Radio del Tanque (m)", value=10.0, step=0.5, min_value=1.0)
        with col2:
            altura_tanque = st.number_input("Altura Útil Máxima (m)", value=1.50, step=0.1, min_value=0.5)
        with col3:
            caudal_concesion = st.number_input("Caudal Concesión Constante (L/s)", value=0.0, step=0.1, min_value=0.0)
        
        area_fija_espejo = math.pi * (radio_tanque**2)
        vol_max_sistema = area_fija_espejo * altura_tanque

        st.subheader("2. Cosecha de Aguas Lluvias (Opción 2)")
        habilitar_cosecha = st.checkbox("¿Implementar cosecha de aguas lluvias mediante cubierta?", value=False, key="check_cosecha")
        
        if habilitar_cosecha:
            col4, col5, col6 = st.columns(3)
            with col4:
                largo_tejado = st.number_input("Largo Tejado (m)", value=10.0, step=1.0, min_value=0.0)
            with col5:
                ancho_tejado = st.number_input("Ancho Tejado (m)", value=10.0, step=1.0, min_value=0.0)
            with col6:
                coef_escorrentia = st.number_input("Coeficiente de Escorrentía", value=0.90, step=0.05, min_value=0.0, max_value=1.0)
        else:
            largo_tejado, ancho_tejado, coef_escorrentia = 0.0, 0.0, 0.0

    else:
        st.subheader("1. Diseño de Reservorio Excavado")
        st.info("Se deshabilitan dimensiones de tanques comerciales para usar batimetría de campo.")
        es_excavado = True
        
        col1, col2 = st.columns(2)
        with col1:
            prof_max = st.number_input("Profundidad máxima (m)", value=2.0, step=0.25, min_value=0.5)
        with col2:
            caudal_concesion = st.number_input("Caudal Concesión (L/s)", value=0.0, step=0.1, min_value=0.0)

        # Generación de tabla de batimetría
        intervalos = np.arange(0, prof_max + 0.25, 0.25)
        df_bat_init = pd.DataFrame({
            "Altura (m)": intervalos,
            "Área Espejo (m2)": [0.0] * len(intervalos),
            "Volumen Acumulado (m3)": [0.0] * len(intervalos)
        })
        
        st.write("Ingrese los datos de la batimetría proyectada:")
        df_bat_usuario = st.data_editor(df_bat_init, num_rows="fixed", use_container_width=True)
        
        if df_bat_usuario["Volumen Acumulado (m3)"].max() > 0:
            vol_max_sistema = df_bat_usuario["Volumen Acumulado (m3)"].max()
            # Lógica interna para interpolación
            h_vals = df_bat_usuario["Altura (m)"].values
            a_vals = df_bat_usuario["Área Espejo (m2)"].values
            v_vals = df_bat_usuario["Volumen Acumulado (m3)"].values
            
            # Funciones de apoyo para la simulación
            func_area = lambda v: np.interp(v, v_vals, a_vals)
        else:
            st.warning("⚠️ Complete la tabla de batimetría para habilitar la simulación.")

    if st.button("Simular Tránsito del Reservorio", type="primary"):
        if 'df_chrono' not in st.session_state or 'q_diseno_decadal' not in st.session_state:
            st.warning("⚠️ Por favor, ejecuta primero el cálculo en la 'Pestaña 2' para generar las matrices de demanda.")
        else:
            with st.spinner("Simulando balance volumétrico y calculando optimización agronómica... 🌊"):
                df_chrono = st.session_state['df_chrono'].copy()
                q_diseno_decadal = st.session_state['q_diseno_decadal']
                t_max = st.session_state.get('t_max', 12)
                area_cultivo_ha = st.session_state.get('area_total_ha', 0.5)
                tipo_riego = st.session_state.get('tipo_riego', "Riego por goteo")
                
                # ---------------------------------------------------------
                # 1. GEOMETRÍA INICIAL (AQUÍ ENTRA EL IF/ELSE DEL TIPO DE RESERVORIO)
                # ---------------------------------------------------------
                if es_excavado:
                    v_max = vol_max_sistema # Viene de la tabla batimétrica (arriba)
                    area_tejado_efectiva = 0.0 # Se asume que no hay cosecha de techos para el reservorio
                else:
                    area_tanque = math.pi * (radio_tanque ** 2)
                    v_max = area_tanque * altura_tanque
                    area_tejado_efectiva = (largo_tejado * ancho_tejado) * coef_escorrentia
                
                dias_d = np.array([10,10,11, 10,10,8, 10,10,11, 10,10,10, 10,10,11, 10,10,10, 10,10,11, 10,10,11, 10,10,10, 10,10,11, 10,10,10, 10,10,11])
                
                resultados_simulacion = []
                v_actual = v_max  # Inicia lleno
                deficit_maximo_registrado = 0.0 
                
                for index, row in df_chrono.iterrows():
                    año, decada_año = int(row['Año']), int(row['Decada_Año'])
                    decada_idx = decada_año - 1 
                    dias, p_dec_mm, e_dec_mm = dias_d[decada_idx], row['Precipitacion'], row['Evaporacion']
                    
                    # ---------------------------------------------------------
                    # 2. ÁREA DINÁMICA DE EVAPORACIÓN/LLUVIA
                    # ---------------------------------------------------------
                    if es_excavado:
                        # Evalúa el polinomio para sacar el área según el volumen que nos queda
                        area_espejo_actual = func_area(v_actual) if v_actual > 0 else func_area(0)
                    else:
                        area_espejo_actual = area_tanque

                    # ENTRADAS
                    e_cp = (caudal_concesion * 86400 * dias) / 1000.0
                    e_ll = area_espejo_actual * (p_dec_mm / 1000.0) # Lluvia directa usa el área dinámica
                    e_es = area_tejado_efectiva * (p_dec_mm / 1000.0) if not es_excavado else 0.0
                    
                    # SALIDAS
                    if tipo_riego == "Riego por goteo":
                        s_d = (q_diseno_decadal[decada_idx] * t_max * 3600 * dias) / 1000.0
                    else:
                        s_d = (q_diseno_decadal[decada_idx] * 86.4 * dias) 
                    
                    s_e = area_espejo_actual * (e_dec_mm / 1000.0) # Evaporación usa el área dinámica
                    s_i = s_e * 0.10 # Asumes infiltración como 10% de evaporación
                    
                    # BALANCE
                    v_temp = v_actual + e_cp + e_ll + e_es - s_d - s_e - s_i
                    derramado, deficit_decada = 0.0, 0.0
                    
                    if v_temp > v_max:
                        v_final, derramado, estado = v_max, v_temp - v_max, "Lleno (Derrama)"
                    elif v_temp < 0:
                        v_final, deficit_decada, estado = 0.0, abs(v_temp), "Déficit Crítico ⚠️"
                        v_actual_matematico = v_actual + e_cp + e_ll + e_es - s_d - s_e - s_i
                        if abs(v_actual_matematico) > deficit_maximo_registrado: 
                            deficit_maximo_registrado = abs(v_actual_matematico)
                    else:
                        v_final, estado = v_temp, "Operación Normal"
                        
                    # ---------------------------------------------------------
                    # 3. ALTURA DE LÁMINA FINAL DE LA DÉCADA
                    # ---------------------------------------------------------
                    if es_excavado:
                        # Si la UI lo definió, interpolamos. (h_vals y v_vals se definieron al llenar la tabla)
                        altura_vaso = np.interp(v_final, v_vals, h_vals) if v_final > 0 else 0
                    else:
                        altura_vaso = v_final / area_tanque if area_tanque > 0 else 0
                    
                    resultados_simulacion.append({
                        'Año': año, 'Decada': decada_año, 'Altura Vaso (m)': round(altura_vaso, 2), 'Volumen Inicial (m3)': round(v_actual, 2),
                        'Entrada Concesion (m3)': round(e_cp, 2), 'Entrada Lluvia (m3)': round(e_ll, 2), 'Entrada Escorrentia (m3)': round(e_es, 2),
                        'Salida Riego (m3)': round(s_d, 2), 'Salida Evaporación (m3)': round(s_e, 2), 'Salida Infiltración (m3)': round(s_i, 2),
                        'Volumen Final (m3)': round(v_final, 2), 'Déficit Hídrico (m3)': round(deficit_decada, 2), 'Volumen Derramado (m3)': round(derramado, 2), 'Estado': estado
                    })
                    v_actual = v_final
                
                df_simulacion = pd.DataFrame(resultados_simulacion)
                st.session_state['df_simulacion_reservorio'] = df_simulacion # Lo guardamos para exportarlo al Word después
                st.success("✅ Simulación de tránsito del reservorio finalizada.")
                
                # Mostrar resultados rápido para el usuario
                st.dataframe(df_simulacion)

               # --- PLANITO ESQUEMÁTICO A ESCALA REAL ---
                st.divider()
                st.subheader("🗺️ Esquema Espacial del Proyecto (Vista en Planta)")
                
                import plotly.graph_objects as go
                fig_esq = go.Figure()

                # --- 1. DEFINIR VARIABLES BASE PRIMERO ---
                # Traemos el área del cultivo (Si no existe, asumimos 0.5 hectáreas por defecto)
                area_cultivo_ha = st.session_state.get('area_total_ha', 0.5)
                area_cultivo_m2 = area_cultivo_ha * 10000
                lado_cultivo = math.sqrt(area_cultivo_m2)
                margen = 5.0 # Margen de separación entre obras en metros
                
                # --- 2. DEFINIR DIMENSIONES DEL ALMACENAMIENTO ---
                if es_excavado:
                    # Calcular el área máxima evaluando el polinomio en el volumen máximo
                    area_maxima = func_area(vol_max_sistema) if vol_max_sistema > 0 else 100.0
                    
                    # Asumimos una forma cuadrada para el espejo de agua máximo
                    lado_reservorio = math.sqrt(area_maxima)
                    distancia_centro = lado_reservorio / 2
                    
                    xc_tanque = lado_cultivo + margen + distancia_centro
                    yc_tanque = margen + distancia_centro
                    forma_dibujo = "rectangulo"
                    
                else:
                    # Usar el radio del tanque australiano (con valor seguro por defecto)
                    distancia_centro = locals().get('radio_tanque', 5.0)
                    xc_tanque = lado_cultivo + margen + distancia_centro
                    yc_tanque = margen + distancia_centro
                    forma_dibujo = "circulo"
                
                # --- 3. DIBUJAR EL CULTIVO ---
                fig_esq.add_shape(
                    type="rect",
                    x0=0, y0=0, x1=lado_cultivo, y1=lado_cultivo,
                    line_color="DarkGreen", fillcolor="LightGreen", opacity=0.3
                )
                fig_esq.add_annotation(x=lado_cultivo/2, y=lado_cultivo/2, text=f"Área de Cultivo<br>({area_cultivo_ha} ha)", showarrow=False)

                # --- 4. DIBUJAR EL ALMACENAMIENTO ---
                if forma_dibujo == "circulo":
                    fig_esq.add_shape(
                        type="circle",
                        x0=xc_tanque - distancia_centro, y0=yc_tanque - distancia_centro,
                        x1=xc_tanque + distancia_centro, y1=yc_tanque + distancia_centro,
                        line_color="DarkBlue", fillcolor="LightSkyBlue"
                    )
                    fig_esq.add_annotation(x=xc_tanque, y=yc_tanque, text="Tanque<br>Australiano", showarrow=False)
                
                elif forma_dibujo == "rectangulo":
                    fig_esq.add_shape(
                        type="rect",
                        x0=xc_tanque - distancia_centro, y0=yc_tanque - distancia_centro,
                        x1=xc_tanque + distancia_centro, y1=yc_tanque + distancia_centro,
                        line_color="SaddleBrown", fillcolor="MediumTurquoise",
                        opacity=0.8
                    )
                    fig_esq.add_annotation(x=xc_tanque, y=yc_tanque, text="Reservorio<br>Excavado", showarrow=False)

                # --- 5. DIBUJAR COSECHA DE AGUAS LLUVIAS (Si aplica) ---
                area_tejado_fisica = locals().get('area_tejado_fisica', 0)
                habilitar_cosecha = locals().get('habilitar_cosecha', False)

                if habilitar_cosecha and area_tejado_fisica > 0:
                    largo_tejado = locals().get('largo_tejado', 10)
                    ancho_tejado = locals().get('ancho_tejado', 10)
                    y_tej_base = yc_tanque + distancia_centro + margen
                    
                    fig_esq.add_shape(
                        type="rect", 
                        x0=lado_cultivo + margen, y0=y_tej_base, 
                        x1=lado_cultivo + margen + largo_tejado, y1=y_tej_base + ancho_tejado, 
                        line=dict(color="DimGray", width=2), fillcolor="rgba(169,169,169,0.6)"
                    )
                    fig_esq.add_annotation(
                        x=lado_cultivo + margen + (largo_tejado/2), 
                        y=y_tej_base + (ancho_tejado/2), 
                        text=f"Cubierta<br>({area_tejado_fisica:,.0f} m²)", showarrow=False
                    )

                # --- 6. CONFIGURACIÓN Y RENDERIZADO DEL GRÁFICO ---
                fig_esq.update_layout(
                    xaxis=dict(scaleanchor="y", scaleratio=1, showgrid=False, zeroline=False, visible=False),
                    yaxis=dict(showgrid=False, zeroline=False, visible=False),
                    plot_bgcolor="white", margin=dict(l=0, r=0, t=30, b=0),
                    title_text="Distribución Espacial a Escala Real", title_x=0.5
                )
                
                st.plotly_chart(fig_esq, use_container_width=True)
                st.caption("🔍 *Nota: Este plano geométrico está renderizado a escala real (1:1). Ayuda a dimensionar la magnitud de las obras civiles frente a la extensión agrícola.*")

                # --- ANÁLISIS DE RESILIENCIA Y AUTO-DIMENSIONAMIENTO ---
                st.divider()
                st.subheader("🛠️ Diagnóstico de Resiliencia y Auto-Dimensionamiento")
                
                col_diag1, col_diag2 = st.columns(2)
                
                with col_diag1:
                    # 1. Resumen seguro (evita el NameError de radio_tanque)
                    if es_excavado:
                        texto_almacenamiento = f"* Tipo: Reservorio Excavado\n* Volumen Máximo: {v_max:.2f} m³"
                    else:
                        radio_seguro = locals().get('radio_tanque', 0.0)
                        texto_almacenamiento = f"* Tipo: Tanque Australiano\n* Radio: {radio_seguro} m\n* Volumen: {v_max:.2f} m³"

                    # Validar si existe cosecha de techos
                    habilitar_cosecha = locals().get('habilitar_cosecha', False)
                    area_tejado_fisica = locals().get('area_tejado_fisica', 0.0)
                    area_tejado_texto = f"{area_tejado_fisica:.2f} m²" if habilitar_cosecha and area_tejado_fisica > 0 else "No implementada"
                    area_cultivo_ha = st.session_state.get('area_total_ha', 0.5)

                    st.info(f"**Diseño Actual Analizado:**\n{texto_almacenamiento}\n* Área Cultivo: {area_cultivo_ha:.2f} Ha\n* Área Ramada: {area_tejado_texto}")
                
                with col_diag2:
                    if deficit_maximo_registrado > 0:
                        st.error(f"🚨 **ALERTA DE QUIEBRE:** El sistema colapsó en época seca. Faltaron hasta **{deficit_maximo_registrado:.2f} m³** de agua en el peor momento.")
                        
                        # 1. Ajuste Estructural
                        volumen_ideal = v_max + deficit_maximo_registrado
                        
                        # 2. Ajuste Agronómico (Búsqueda Binaria del Área Óptima)
                        low, high = 0.001, area_cultivo_ha
                        area_optima = 0.0
                        
                        for _ in range(20): # 20 iteraciones (precisión de 0.001 Ha)
                            mid = (low + high) / 2
                            factor_area = mid / area_cultivo_ha
                            v_act_sim = v_max
                            fallo_sim = False
                            
                            for idx_s, row_s in df_chrono.iterrows():
                                d_idx = int(row_s['Decada_Año']) - 1
                                p_mm, e_mm = row_s['Precipitacion'], row_s['Evaporacion']
                                
                                # --- Lógica dual: Excavado vs Tanque ---
                                if es_excavado:
                                    area_sim = func_area(v_act_sim) if v_act_sim > 0 else func_area(0)
                                    e_es_s = 0.0 # No se asume cosecha de techos para reservorio excavado
                                else:
                                    area_sim = locals().get('area_tanque', 0.0)
                                    area_tej_ef = locals().get('area_tejado_efectiva', 0.0)
                                    e_es_s = area_tej_ef * (p_mm / 1000.0)
                                
                                e_cp_s = (caudal_concesion * 86400 * dias_d[d_idx]) / 1000.0
                                e_ll_s = area_sim * (p_mm / 1000.0)
                                
                                # Salida de riego ajustada por el factor de iteración
                                if tipo_riego == "Riego por goteo":
                                    s_d_s = (q_diseno_decadal[d_idx] * factor_area * t_max * 3600 * dias_d[d_idx]) / 1000.0
                                else:
                                    s_d_s = (q_diseno_decadal[d_idx] * factor_area * 86.4 * dias_d[d_idx])
                                
                                s_e_s = area_sim * (e_mm / 1000.0)
                                s_i_s = s_e_s * 0.10
                                
                                v_temp_s = v_act_sim + e_cp_s + e_ll_s + e_es_s - s_d_s - s_e_s - s_i_s
                                
                                if v_temp_s < 0:
                                    fallo_sim = True
                                    break
                                v_act_sim = v_max if v_temp_s > v_max else v_temp_s
                                
                            if fallo_sim:
                                high = mid # El área sigue siendo muy grande
                            else:
                                low = mid  # El área soporta bien, intentemos un poco más
                                area_optima = mid
                        
                        # Mostrar Soluciones (Diferenciadas por infraestructura)
                        if es_excavado:
                            st.info(f"🏗️ **Opción 1 (Estructural):** Se requiere rediseñar la topografía/batimetría del reservorio para alcanzar un volumen útil de al menos **{volumen_ideal:.2f} m³**.")
                        else:
                            altura_segura = locals().get('altura_tanque', 1.5)
                            radio_ideal = math.sqrt(volumen_ideal / (math.pi * altura_segura))
                            
                            if radio_ideal > 20.0:
                                st.warning(f"🏗️ **Opción 1 (Estructural):** Subir el radio a **{radio_ideal:.2f} m** es inviable (>20m). Te recomendamos aumentar el área de la ramada.")
                            else:
                                st.info(f"🏗️ **Opción 1 (Estructural):** Incrementa el radio del tanque a **{radio_ideal:.2f} metros** para mantener las {area_cultivo_ha:.2f} Ha actuales de cultivo.")
                            
                        if area_optima > 0.01:
                            st.success(f"🌱 **Opción 2 (Agronómica):** Si no puedes agrandar el reservorio, debes reducir el área de cultivo a máximo **{area_optima:.3f} Ha** ({(area_optima*10000):.0f} m²) para garantizar agua todo el año.")
                        else:
                            st.error(f"🌱 **Opción 2 (Agronómica):** El reservorio propuesto es tan pequeño que no puede sostener ni 0.01 Ha. ¡Necesitas rediseñar urgentemente o buscar una concesión de agua!")
                    else:
                        st.success("🏆 **DISEÑO ÓPTIMO:** El reservorio propuesto es resiliente y no se vació durante toda la serie climática analizada.")
                


# --- PESTAÑA 4: GENERACIÓN DE MEMORIA DE CÁLCULO ---

with tab4:
    st.header("Generación de Memorias de Cálculo")
    col_a, col_b = st.columns(2)
    
    with col_a:
        if st.button("Generar Anexo 3 (Hidrología)"):
            # 1. Recuperar los datos de la sesión guardados en las pestañas previas
            df_clima = st.session_state.get('df_chrono', None)
            df_sim = st.session_state.get('df_simulacion_reservorio', None)
            
            # (Opcional) Guardar en Session State el tipo de reservorio y volumen maximo en la Pestaña 3 
            # para llamarlos aquí. Si no los tienes, puedes poner valores por defecto:
            tipo_almacenamiento = st.session_state.get('tipo_almacenamiento_elegido', "Reservorio Excavado")
            vol_maximo = st.session_state.get('volumen_maximo_sistema', 0.0)

            # 2. Verificar que se haya hecho la simulación
            if df_sim is None:
                st.error("⚠️ Debes ir a la Pestaña 3 y dar clic en 'Simular Tránsito del Reservorio' antes de generar este documento.")
            else:
                # 3. Crear documento y descargar
                doc_h = crear_memoria_hidrologia(
                    datos_clima=df_clima, 
                    coordenadas=None, 
                    df_simulacion=df_sim, 
                    tipo_almacenamiento=tipo_almacenamiento, 
                    vol_max=vol_maximo
                )
                
                buffer = io.BytesIO()
                doc_h.save(buffer)
                st.download_button(
                    label="📥 Descargar Anexo 3 (.docx)", 
                    data=buffer.getvalue(), 
                    file_name="Anexo_3_Hidrologia_ADR.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

    with col_b:
        if st.button("Generar Memoria de Demandas"):
            doc_d = crear_memoria_demandas(None)
            buffer = io.BytesIO()
            doc_d.save(buffer)
            st.download_button("Descargar Memoria Demandas", buffer.getvalue(), "Memoria_Demandas_Riego.docx")

